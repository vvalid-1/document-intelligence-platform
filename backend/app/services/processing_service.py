from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

import chardet
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from pypdf import PdfReader

logger = logging.getLogger(__name__)

_OCR_CHAR_THRESHOLD = 10
_OCR_DPI = 200
_OCR_TIMEOUT = 30


@dataclass
class PageText:
    page_number: Optional[int]  # 1-indexed for PDFs; None for DOCX/TXT
    text: str


@dataclass
class Chunk:
    chunk_index: int
    chunk_text: str
    page_number: Optional[int]
    token_count: int  # word-count approximation


# ── Extraction ────────────────────────────────────────────────────────────────

def extract_text_pdf(path: Path) -> list[PageText]:
    pages: list[PageText] = []
    with fitz.open(str(path)) as pdf:
        for i in range(len(pdf)):
            page = pdf[i]
            text = page.get_text().strip()
            if len(text) < _OCR_CHAR_THRESHOLD:
                text = _ocr_page(page)
            pages.append(PageText(page_number=i + 1, text=text))
    return pages


def _ocr_page(page: fitz.Page) -> str:
    try:
        pix = page.get_pixmap(dpi=_OCR_DPI)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        return pytesseract.image_to_string(img, timeout=_OCR_TIMEOUT).strip()
    except Exception as exc:
        logger.warning("OCR failed on page %d: %s", page.number + 1, exc)
        return ""


def extract_text_docx(path: Path) -> str:
    import docx  # python-docx
    document = docx.Document(str(path))
    parts: list[str] = []
    for para in document.paragraphs:
        stripped = para.text.strip()
        if stripped:
            parts.append(stripped)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_text_image(path: Path) -> str:
    try:
        img = Image.open(str(path))
        return pytesseract.image_to_string(img, timeout=_OCR_TIMEOUT).strip()
    except Exception as exc:
        logger.warning("OCR failed on image %s: %s", path.name, exc)
        return ""


def extract_text_txt(path: Path) -> str:
    raw = path.read_bytes()
    detected = chardet.detect(raw)
    encoding: str = detected.get("encoding") or "utf-8"
    return raw.decode(encoding, errors="replace")


def extract_pdf_metadata(path: Path) -> dict:
    reader = PdfReader(str(path))
    pdf_meta = reader.metadata or {}
    result: dict = {"page_count": len(reader.pages)}
    for key, label in (
        ("/Author", "author"),
        ("/Title", "title"),
        ("/Subject", "subject"),
        ("/CreationDate", "creation_date"),
    ):
        val = pdf_meta.get(key)
        if val:
            result[label] = str(val)
    return result


# ── Chunking ──────────────────────────────────────────────────────────────────

def chunk_pages(pages: list[PageText], chunk_size: int, overlap: int) -> list[Chunk]:
    """Chunk PDF pages into overlapping windows, preserving page attribution."""
    segments = [(p.page_number, p.text) for p in pages if p.text.strip()]
    if not segments:
        return []

    # Build a flat character list with per-char page tracking
    flat_chars: list[str] = []
    page_map: list[Optional[int]] = []
    sep = "\n\n"

    for i, (page_num, text) in enumerate(segments):
        for ch in text:
            flat_chars.append(ch)
            page_map.append(page_num)
        if i < len(segments) - 1:
            for ch in sep:
                flat_chars.append(ch)
                page_map.append(page_num)

    full_text = "".join(flat_chars)
    return _chunk(full_text, page_map, chunk_size, overlap)


def chunk_plain_text(text: str, chunk_size: int, overlap: int) -> list[Chunk]:
    """Chunk DOCX/TXT text with no page tracking."""
    stripped = text.strip()
    if not stripped:
        return []
    return _chunk(stripped, [], chunk_size, overlap)


def _chunk(
    text: str,
    page_map: list[Optional[int]],
    chunk_size: int,
    overlap: int,
) -> list[Chunk]:
    if overlap >= chunk_size:
        overlap = chunk_size // 4

    chunks: list[Chunk] = []
    idx = 0
    start = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))
        window = text[start:end]
        chunk_text = window.strip()

        if chunk_text:
            leading_spaces = len(window) - len(window.lstrip())
            actual_start = min(start + leading_spaces, len(page_map) - 1) if page_map else 0
            page_num: Optional[int] = page_map[actual_start] if page_map else None

            chunks.append(
                Chunk(
                    chunk_index=idx,
                    chunk_text=chunk_text,
                    page_number=page_num,
                    token_count=len(chunk_text.split()),
                )
            )
            idx += 1

        if end == len(text):
            break
        start += chunk_size - overlap

    return chunks
